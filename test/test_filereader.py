import pytest
from unittest.mock import patch, mock_open
from io import BytesIO
from fastapi import UploadFile
from docx import Document
import time

from FileReader import (
    read_docx_file,
    split_document_into_sections,
    extract_exercises,
    read_file,
    read_all_files,
    extract_title,
    extract_author
)

def mock_upload_file(filename: str, content: bytes) -> UploadFile:
    file_mock = UploadFile(filename=filename, file=BytesIO(content))
    return file_mock

# Unit tests
def test_read_docx_file():
    docx_content = BytesIO()
    doc = Document()
    doc.add_paragraph("Paragraph 1")
    doc.add_paragraph("Paragraph 2")
    doc.save(docx_content)
    docx_content.seek(0)

    with patch("builtins.open", mock_open(read_data=docx_content.read())):
        paragraphs = read_docx_file(docx_content)
        assert paragraphs == ["Paragraph 1", "Paragraph 2"]


def test_split_document_into_sections():
    content = [
        "Title: Test Title",
        "Author: Test Author",
        "1. Experiment aim:",
        "This is the aim.",
        "2. Theoretical background:",
        "This is the background.",
        "3. Research:",
        "Ex. 1. This is exercise 1.",
        "Ex. 2. This is exercise 2.",
        "4. Conclusions:",
        "This is the conclusion."
    ]
    patterns = ["Title:", "Author:", "1. Experiment aim:", "2. Theoretical background:", "3. Research:", "4. Conclusions:"]
    sections = split_document_into_sections(content, patterns)
    assert "Title:" in sections
    assert "Author:" in sections
    assert "1. Experiment aim:" in sections
    assert "2. Theoretical background:" in sections
    assert "3. Research:" in sections
    assert "4. Conclusions:" in sections


def test_extract_exercises():
    sections = {
        "3. Research:": [
            "Ex. 1. This is exercise 1.",
            "Some description.",
            "Ex. 2. This is exercise 2.",
            "Some other description."
        ]
    }
    exercises = extract_exercises(sections, 2)
    assert "Ex. 1." in exercises
    assert "Ex. 2." in exercises

# Integration tests
@pytest.mark.asyncio
async def test_extract_title():
    docx_content1 = BytesIO()
    doc1 = Document()
    doc1.add_paragraph("Title: Test Title")
    doc1.add_paragraph("Content of file 1.")
    doc1.save(docx_content1)
    docx_content1.seek(0)

    docx_content2 = BytesIO()
    doc2 = Document()
    doc2.add_paragraph("Title: Test Title")
    doc2.add_paragraph("Content of file 2.")
    doc2.save(docx_content2)
    docx_content2.seek(0)

    files = [
        mock_upload_file("file1.docx", docx_content1.getvalue()),
        mock_upload_file("file2.docx", docx_content2.getvalue())
    ]
    title, contents = await extract_title(files)
    assert title == "Test Title"
    assert len(contents) == 2

def test_extract_author():
    docx_content1 = BytesIO()
    doc1 = Document()
    doc1.add_paragraph("Author: John Doe 123456")
    doc1.add_paragraph("Content of file 1.")
    doc1.save(docx_content1)
    docx_content1.seek(0)

    docx_content2 = BytesIO()
    doc2 = Document()
    doc2.add_paragraph("Author: Jane Smith 654321")
    doc2.add_paragraph("Content of file 2.")
    doc2.save(docx_content2)
    docx_content2.seek(0)

    contents = [docx_content1.getvalue(), docx_content2.getvalue()]

    with patch("FileReader.add_student") as mock_add_student:
        author_ids = extract_author(contents)
        assert author_ids == ["123456", "654321"]
        assert mock_add_student.call_count == 2

def test_read_file():
    file_path = "dummy_path.docx"
    patterns = ["Title:", "Author:", "1. Experiment aim:", "2. Theoretical background:", "3. Research:", "4. Conclusions:"]
    number_of_exercises = 2
    file_id = 0

    content = [
        "Title: Test Title",
        "Author: Test Author",
        "1. Experiment aim:",
        "This is the aim.",
        "2. Theoretical background:",
        "This is the background.",
        "3. Research:",
        "Ex. 1. This is exercise 1.",
        "Ex. 2. This is exercise 2.",
        "4. Conclusions:",
        "This is the conclusion."
    ]
    docx_content = BytesIO()
    doc = Document()
    for line in content:
        doc.add_paragraph(line)
    doc.save(docx_content)
    docx_content.seek(0)

    with patch("builtins.open", mock_open(read_data=docx_content.read())):
        with patch("FileReader.read_docx_file", return_value=content):
            documents, metadatas = read_file(file_path, patterns, number_of_exercises, file_id)
            assert len(documents) > 0
            assert len(metadatas) > 0

def test_read_all_files():
    folder_path = "dummy_folder"
    patterns = ["Title:", "Author:", "1. Experiment aim:", "2. Theoretical background:", "3. Research:", "4. Conclusions:"]
    number_of_exercises = 2

    files = ["file1.docx", "file2.docx"]
    with patch("os.listdir", return_value=files):
        with patch("FileReader.read_file", return_value=(["doc"], ["meta"])):
            documents, metadatas = read_all_files(folder_path, patterns, number_of_exercises)
            assert len(documents) == 2
            assert len(metadatas) == 2

# performance tests
def test_performance_read_all_files():
    folder_path = "dummy_folder"
    patterns = ["Title:", "Author:", "1. Experiment aim:", "2. Theoretical background:", "3. Research:", "4. Conclusions:"]
    number_of_exercises = 2

    files = [f"file{i}.docx" for i in range(100)]
    with patch("os.listdir", return_value=files):
        with patch("FileReader.read_file", return_value=(["doc"], ["meta"])):
            start_time = time.time()
            documents, metadatas = read_all_files(folder_path, patterns, number_of_exercises)
            end_time = time.time()
            duration = end_time - start_time
            assert len(documents) == 100
            assert len(metadatas) == 100
            assert duration < 5

