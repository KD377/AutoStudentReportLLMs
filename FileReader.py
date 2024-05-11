import os
from fastapi import UploadFile, File, HTTPException
from docx import Document
from typing import List
import re
import io


def read_docx_file(file_path):
    doc = Document(file_path)
    content = []
    for paragraph in doc.paragraphs:
        content.append(paragraph.text)
    return content


def split_document_into_sections(content, patterns):
    sections = {}
    current_section = None

    for line in content:
        for pattern in patterns:
            if line.strip().startswith(pattern):
                current_section = pattern
                sections[current_section] = []
                if not line.startswith("Title:") and not line.startswith("Author:"):
                    sections[current_section].append(current_section)
                if current_section == "Title:":
                    sections[current_section].append(line.strip())
                elif current_section == "Author:":
                    sections[current_section].append(line.strip())
                break
        else:
            if current_section is not None and line.strip():
                sections[current_section].append(line.strip())
    return sections


def remove_non_ascii(input_string):
    cleaned_string = ''.join(char for char in input_string if ord(char) < 128)
    return cleaned_string


def extract_exercises(sections, exercise_number):
    exercises = {}
    current_section = None

    patterns = [f"Ex. {i}." for i in range(1, exercise_number + 1)]
    for line in sections["3. Research:"]:
        for pattern in patterns:
            if line.strip().startswith(pattern):
                current_section = pattern
                exercises[current_section] = []
        else:
            if current_section is not None:
                exercises[current_section].append(line.strip())
    return exercises


def read_file(file_path, patterns, number_of_exercises, file_id):
    sections = split_document_into_sections(read_docx_file(file_path), patterns)
    author = sections["Author:"][0]
    exercises = extract_exercises(sections, number_of_exercises)
    del sections["3. Research:"]
    updated_sections = {}
    ordered_keys = ['Author:', 'Title:', '1. Experiment aim:', '2. Theoretical background:'] + sorted(
        exercises.keys()) + ['4. Conclusions:']
    for key in ordered_keys:
        if key in sections:
            updated_sections[key] = sections[key]
        if key in exercises and key not in sections:
            updated_sections[key] = exercises[key]
    metadatas = []
    pattern = r'Ex\. (?:[1-9]|[1-9][0-9])\.'

    documents = []
    student_answer = False
    i, j = 0, 0
    current_section = None
    for section_name, sentences in updated_sections.items():
        j = 0
        if section_name != current_section:
            student_answer = False
            current_section = section_name
        for sentence in sentences:
            cleaned_sentence = remove_non_ascii(sentence)
            documents.append(cleaned_sentence)
            meta = {"File_ID": file_id}  # Include file ID in metadata

            if sentence.strip().startswith("Student’s answer:"):
                student_answer = True

            if re.findall(pattern, section_name):
                meta["Section_name"] = "3. Research:"
                parts = section_name.split(".")
                if len(parts) == 3 and parts[0] == "Ex" and parts[1].strip().isdigit():
                    number = int(parts[1].strip())
                    meta["Exercise_number"] = number
            else:
                meta["Exercise_number"] = 0
                meta["Section_name"] = section_name

            if student_answer:
                meta["Type"] = "answer"
            else:
                meta["Type"] = "description"
            meta["Global_sentence_number"] = i
            meta["Local_sentence_number"] = j
            meta["Author"] = author[len("Author: "):].strip()
            metadatas.append(meta)

            j += 1
            i += 1
    return documents, metadatas


def read_all_files(folder_path, patterns, number_of_exercises):
    all_documents = []
    all_metadatas = []
    files = [f for f in os.listdir(folder_path) if f.endswith('.docx')]

    file_id = 0
    for file in files:
        file_path = os.path.join(folder_path, file)
        document, metadata = read_file(file_path, patterns, number_of_exercises, file_id)
        all_documents.append(document)
        all_metadatas.append(metadata)
        file_id += 1

    return all_documents, all_metadatas


async def extract_title(files: List[UploadFile] = File(...)):
    titles = []
    contents = []
    for file in files:
        content = await file.read()
        contents.append(content)
        document = Document(io.BytesIO(content))
        title = None
        for para in document.paragraphs:
            if para.text.startswith("Title: "):
                title = para.text[len("Title: "):]
                break

        if not title:
            raise HTTPException(status_code=400, detail=f"No title found in {file.filename}")
        titles.append(title)

    if len(set(titles)) > 1:
        raise HTTPException(status_code=400, detail="Titles are not consistent across all files.")
    return titles[0], contents
